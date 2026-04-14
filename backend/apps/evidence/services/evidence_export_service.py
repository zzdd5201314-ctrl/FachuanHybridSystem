"""Business logic services."""

from __future__ import annotations

import io
from pathlib import Path
from typing import TYPE_CHECKING, Any

from django.utils import timezone
from django.utils.translation import gettext_lazy as _
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from docxtpl import DocxTemplate

from apps.core.exceptions import NotFoundError, ValidationException
from apps.documents.services.placeholders.fallback import build_docx_render_context
from apps.evidence.models import EvidenceItem, EvidenceList

if TYPE_CHECKING:
    from apps.core.interfaces import IEvidenceListPlaceholderService


class EvidenceExportService:
    """
    证据导出服务

    负责生成证据清单和证据明细的 Word 文档.
    支持默认导出和模板导出两种方式.

    Requirements: 6.1, 6.2, 6.3, 6.4
    """

    def __init__(self, placeholder_service: IEvidenceListPlaceholderService | None = None) -> None:
        """
        初始化服务

        Args:
            placeholder_service: 占位符服务(可选,支持依赖注入)
        """
        self._placeholder_service = placeholder_service

    @property
    def placeholder_service(self) -> IEvidenceListPlaceholderService:
        """
        获取占位符服务(延迟加载)

        Returns:
            IEvidenceListPlaceholderService 实例
        """
        if self._placeholder_service is None:
            from apps.evidence.services.wiring import get_evidence_list_placeholder_service

            self._placeholder_service = get_evidence_list_placeholder_service()
        return self._placeholder_service

    def export_evidence_list_with_template(self, list_id: int, template_id: int | None = None) -> tuple[bytes, str]:
        """
        使用模板导出证据清单

        Args:
            list_id: 证据清单 ID
            template_id: 模板 ID(可选,为空时使用默认导出)

        Returns:
            (文档内容, 文件名)

        Raises:
            NotFoundError: 证据清单或模板不存在
            ValidationException: 模板渲染失败

        Requirements: 6.1, 6.2, 6.3, 6.4
        """
        # 获取证据清单
        evidence_list = self._get_evidence_list(list_id)

        # 如果没有指定模板,使用默认导出
        # Requirements: 6.2
        if template_id is None:
            return self.export_evidence_list(list_id)

        # 获取模板
        template = self._get_template(template_id)

        # 获取模板文件路径
        template_path = template.get_file_location()
        if not template_path:
            raise NotFoundError(
                message=_("模板文件不存在"),
                code="TEMPLATE_FILE_NOT_FOUND",
                errors={"template": f"模板 {template_id} 的文件路径为空"},
            )

        # 检查模板文件是否存在
        from apps.core.utils.path import Path

        if not Path(template_path).exists():
            raise NotFoundError(
                message=_("模板文件不存在"),
                code="TEMPLATE_FILE_NOT_FOUND",
                errors={"template": f"模板 {template_id} 的文件不存在: {template_path}"},
            )

        # 获取占位符上下文
        # Requirements: 6.3
        try:
            context = self.placeholder_service.get_evidence_list_context(list_id)
        except Exception as e:
            import traceback

            raise ValidationException(
                message=_("获取占位符上下文失败"),
                code="TEMPLATE_RENDER_ERROR",
                errors={"context": f"获取占位符数据时发生错误: {e!s}\n{traceback.format_exc()}"},
            ) from e

        # 使用 docxtpl 渲染模板
        # Requirements: 6.1
        try:
            doc = DocxTemplate(template_path)
            doc.render(build_docx_render_context(doc=doc, context=context))

            # 保存到内存
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            content = buffer.read()
        except Exception as e:
            # Requirements: 6.4
            import traceback

            raise ValidationException(
                message=_("模板渲染失败"),
                code="TEMPLATE_RENDER_ERROR",
                errors={"template": f"渲染模板时发生错误: {e!s}\n{traceback.format_exc()}"},
            ) from e

        # 获取版本号
        version = self._increment_version(evidence_list)

        # 生成文件名
        filename = self._generate_filename(evidence_list, "证据清单", version)

        return content, filename

    def _get_template(self, template_id: int) -> Any:
        """
        获取文档模板

        Args: template_id: 模板 ID

        Returns:
            DocumentTemplate 实例

        Raises:
            NotFoundError: 模板不存在
        """
        from apps.evidence.models import DocumentCaseFileSubType, DocumentTemplate, DocumentTemplateType

        try:
            return DocumentTemplate.objects.get(
                id=template_id,
                is_active=True,
                template_type=DocumentTemplateType.CASE,
                case_sub_type=DocumentCaseFileSubType.EVIDENCE_MATERIALS,
            )
        except DocumentTemplate.DoesNotExist:
            raise NotFoundError(
                message=_("模板不存在"),
                code="TEMPLATE_NOT_FOUND",
                errors={"template_id": f"ID 为 {template_id} 的模板不存在或已禁用"},
            ) from None

    def export_evidence_list(self, list_id: int) -> tuple[bytes, str]:
        """
        导出证据清单为 Word 文档

        Args:
            list_id: 证据清单 ID

        Returns:
            (文档内容, 文件名)

        Raises:
            NotFoundError: 证据清单不存在

        Requirements: 8.1-8.10
        """
        evidence_list = self._get_evidence_list(list_id)
        items = list(evidence_list.items.order_by("order"))

        # 计算全局序号起始值(前面所有清单的证据数量 + 1)
        global_order_start = self._get_global_order_start(evidence_list)

        # 获取版本号
        version = self._increment_version(evidence_list)

        # 创建文档
        doc = Document()

        # 设置标题
        title = doc.add_heading(evidence_list.title, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加案件信息
        case_info = doc.add_paragraph()
        case_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_info.add_run(f"案件:{evidence_list.case.name}")

        # 添加空行
        doc.add_paragraph()

        # 创建表格(传入全局序号起始值)
        self._create_evidence_table(doc, items, global_order_start)

        # 生成文件名
        filename = self._generate_filename(evidence_list, "证据清单", version)

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read(), filename

    def export_evidence_detail(self, list_id: int) -> tuple[bytes, str]:
        """
        导出证据明细为 Word 文档

        Args:
            list_id: 证据清单 ID

        Returns:
            (文档内容, 文件名)

        Raises:
            NotFoundError: 证据清单不存在

        Requirements: 10.1-10.5
        """
        evidence_list = self._get_evidence_list(list_id)
        items = list(evidence_list.items.order_by("order"))

        # 计算全局序号起始值
        global_order_start = self._get_global_order_start(evidence_list)

        # 获取版本号
        version = self._increment_version(evidence_list)

        # 创建文档
        doc = Document()

        # 设置标题
        title = doc.add_heading("证据明细", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 添加案件信息
        case_info = doc.add_paragraph()
        case_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        case_info.add_run(f"案件:{evidence_list.case.name}")
        case_info.add_run(f"  证据清单:{evidence_list.title}")

        # 添加空行
        doc.add_paragraph()

        # 为每个证据添加详细信息(使用全局序号)
        for index, item in enumerate(items):
            global_order = global_order_start + index
            self._add_evidence_detail_section(doc, item, global_order)

        # 生成文件名
        filename = self._generate_filename(evidence_list, "证据明细", version)

        # 保存到内存
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.read(), filename

    def _get_evidence_list(self, list_id: int) -> Any:
        """获取证据清单"""
        try:
            return EvidenceList.objects.select_related("case").get(id=list_id)
        except EvidenceList.DoesNotExist:
            raise NotFoundError(
                message=_("证据清单不存在"),
                code="EVIDENCE_LIST_NOT_FOUND",
                errors={"list_id": f"ID 为 {list_id} 的证据清单不存在"},
            ) from None

    def _create_evidence_table(self, doc: Document, items: list[EvidenceItem], global_order_start: int = 1) -> None:
        """
        创建证据清单表格

        Args:
            doc: Word 文档
            items: 证据明细列表
            global_order_start: 全局序号起始值

        Requirements: 8.2, 8.3
        """
        # 创建表格:序号、证据名称、证据种类、证明内容、原件状态、页码
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置列宽
        widths = [Cm(1.2), Cm(3.5), Cm(2), Cm(6), Cm(1.8), Cm(1.8)]
        for i, width in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = width

        # 表头
        header_cells = table.rows[0].cells
        headers = ["序号", "证据名称", "证据种类", "证明内容", "原件/复印件", "页码"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in header_cells[i].paragraphs[0].runs:
                run.bold = True

        # 数据行(使用全局序号)
        for index, item in enumerate(items):
            global_order = global_order_start + index
            row_cells = table.add_row().cells
            row_cells[0].text = str(global_order)
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[1].text = item.name
            row_cells[2].text = item.get_evidence_type_display() if item.evidence_type else ""
            row_cells[3].text = item.purpose
            row_cells[4].text = item.get_original_status_display() if item.original_status else ""
            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row_cells[5].text = item.page_range_display
            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_evidence_detail_section(self, doc: Document, item: EvidenceItem, global_order: int) -> None:
        """
        添加证据明细章节

        Args:
            doc: Word 文档
            item: 证据明细
            global_order: 全局序号

        Requirements: 10.2
        """
        # 证据标题(使用全局序号)
        doc.add_heading(f"证据 {global_order}:{item.name}", level=2)
        # 证据种类/方向
        meta_parts: list[str] = []
        if item.direction:
            meta_parts.append(f"方向:{item.get_direction_display()}")
        if item.evidence_type:
            meta_parts.append(f"种类:{item.get_evidence_type_display()}")
        if item.original_status:
            meta_parts.append(f"原件状态:{item.get_original_status_display()}")
        if meta_parts:
            doc.add_paragraph("  ".join(meta_parts))
        # 证明内容
        doc.add_paragraph(f"证明内容:{item.purpose}")
        # 页码范围
        if item.page_start and item.page_end:
            doc.add_paragraph(f"页码范围:{item.page_range_display}")
        # 文件信息
        if item.file:
            doc.add_paragraph(f"文件名:{item.file_name}")
            doc.add_paragraph(f"文件大小:{item.file_size_display}")
            doc.add_paragraph(f"页数:{item.page_count}")
        # 添加分隔线
        doc.add_paragraph("─" * 50)

    def _get_global_order_start(self, evidence_list: EvidenceList) -> int:
        """
        计算全局序号起始值

        全局序号 = 前面所有清单的证据数量 + 1

        Args:
            evidence_list: 当前证据清单

        Returns:
            全局序号起始值
        """
        from django.db.models import Count

        previous_count = (
            EvidenceList.objects.filter(case_id=evidence_list.case_id, order__lt=evidence_list.order).aggregate(
                total=Count("items")
            )["total"]
            or 0
        )

        return previous_count + 1

    def _generate_filename(self, evidence_list: EvidenceList, doc_type: str, version: int) -> str:
        """
        生成导出文件名

        Args:
            evidence_list: 证据清单
            doc_type: 文档类型(证据清单/证据明细)
            version: 版本号

        Returns:
            文件名

        Requirements: 8.7, 10.3

        文件名格式:
        - 证据清单:{证据清单名称}({案件名称})V{版本号}_{日期}.docx
          示例:证据清单一(XX与YY纠纷)V1_20260118.docx
        - 证据明细:证据明细{清单序号}({案件名称})V{版本号}_{日期}.docx
          示例:证据明细一(XX与YY纠纷)V1_20260118.docx
        """
        # 获取案件名称
        case_name = evidence_list.case.name

        # 获取日期(格式:YYYYMMDD)
        date_str = timezone.now().strftime("%Y%m%d")

        # 根据文档类型生成文件名
        if doc_type == "证据清单":
            # 格式:{证据清单名称}({案件名称})V{版本号}_{日期}.docx
            filename = f"{evidence_list.title}({case_name})V{version}_{date_str}.docx"
        else:
            # 格式:证据明细{清单序号}({案件名称})V{版本号}_{日期}.docx
            # 从证据清单标题中提取序号部分(如"证据清单一"提取"一")
            list_suffix = ""
            title = evidence_list.title
            if title.startswith("证据清单"):
                list_suffix = title[4:]  # 提取"证据清单"后面的部分
            elif title.startswith("补充证据清单"):
                list_suffix = title[6:]  # 提取"补充证据清单"后面的部分
            filename = f"证据明细{list_suffix}({case_name})V{version}_{date_str}.docx"

        return filename

    def _increment_version(self, evidence_list: EvidenceList) -> Any:
        """
        获取当前导出版本号(不再自动递增)

        版本号由用户在 Admin 界面手动控制.

        Args:
            evidence_list: 证据清单

        Returns:
            当前版本号

        Requirements: 8.8, 8.9, 8.10, 10.4, 10.5
        """
        # 直接返回当前版本号,不再自动递增
        return evidence_list.export_version

    def export_zip(self, list_id: int) -> tuple[bytes, str]:
        """
        导出证据清单为 ZIP 包（含所有证据文件 + Word 清单）

        Returns:
            (ZIP 内容, 文件名)
        """
        import zipfile

        evidence_list = self._get_evidence_list(list_id)
        items = list(evidence_list.items.order_by("order"))
        global_start = self._get_global_order_start(evidence_list)

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            # 添加 Word 清单
            word_content, word_name = self.export_evidence_list(list_id)
            zf.writestr(word_name, word_content)

            # 添加各证据文件
            for idx, item in enumerate(items):
                if not item.file:
                    continue
                order = global_start + idx
                ext = Path(item.file_name).suffix if item.file_name else ""
                arc_name = f"{order:03d}_{item.name}{ext}"
                try:
                    item.file.seek(0)
                    zf.writestr(arc_name, item.file.read())
                except Exception:
                    pass

        buf.seek(0)
        version = evidence_list.export_version
        date_str = timezone.now().strftime("%Y%m%d")
        filename = f"{evidence_list.title}({evidence_list.case.name})V{version}_{date_str}.zip"
        return buf.read(), filename
