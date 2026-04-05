"""Business logic services."""

import logging
import re

logger = logging.getLogger(__name__)


PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([\w\u4e00-\u9fff][\w\u4e00-\u9fff\.\(\)]*)\s*\}\}")


def extract_placeholders(file_path: str) -> list[str]:
    placeholders: set[str] = set()

    try:
        from docxtpl import DocxTemplate

        doc = DocxTemplate(file_path)
        undeclared = doc.get_undeclared_template_variables()
        placeholders.update(undeclared)
    except Exception:
        logger.exception("操作失败")

        pass

    try:
        from docx import Document

        word_doc = Document(file_path)
        for para in word_doc.paragraphs:
            placeholders.update(PLACEHOLDER_PATTERN.findall(para.text))

        for table in word_doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders.update(PLACEHOLDER_PATTERN.findall(cell.text))
    except Exception:
        logger.exception("操作失败")

        pass

    return sorted(list(placeholders))
