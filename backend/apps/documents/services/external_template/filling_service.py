"""
外部模板填充服务

负责占位符取值、填充预览、自定义字段获取、单次模板填充。
批量填充方法将在后续任务 (7.3) 中实现。

Requirements: 6.1-6.10, 10.1-10.6, 11.3, 15.3, 15.4, 15.6,
              16.1-16.6, 17.1-17.3
"""

from __future__ import annotations

import logging
import uuid
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING, Any
from xml.etree import ElementTree as ET

from django.apps import apps
from django.conf import settings
from django.db.models import QuerySet
from django.utils import timezone
from django.utils.translation import gettext_lazy as _

if TYPE_CHECKING:
    from apps.documents.services.placeholders.registry import PlaceholderRegistry

logger: logging.Logger = logging.getLogger(__name__)


@dataclass(frozen=True)
class FillPreviewItem:
    """填充预览项"""

    position_description: str
    semantic_label: str
    fill_value: str
    value_source: str  # "auto" | "manual" | "empty"
    fill_type: str
    mapping_id: int


@dataclass(frozen=True)
class FillReport:
    """填充报告"""

    total_fields: int
    filled_count: int
    skipped_count: int
    manual_needed: list[str]
    errors: list[str]


class FillingService:
    """模板填充服务：占位符取值 + 填充预览 + 自定义字段"""

    def __init__(self, placeholder_registry: PlaceholderRegistry) -> None:
        self._placeholder_registry = placeholder_registry

    # ------------------------------------------------------------------
    # 预览
    # ------------------------------------------------------------------

    def generate_preview(
        self,
        template_id: int,
        case_id: int,
        party_id: int | None = None,
        custom_values: dict[str, str] | None = None,
    ) -> list[FillPreviewItem]:
        """
        生成填充预览：
        1. 获取模板的所有 FieldMapping（按 sort_order 排序）
        2. 从占位符体系获取案件数据（含当事人数据）
        3. 合并自定义值
        4. 返回每个字段的预览信息（位置、语义、值、来源）
        """
        from apps.documents.models.external_template import ExternalTemplateFieldMapping

        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
        ).order_by("sort_order", "id")

        placeholder_values: dict[str, str] = self._get_placeholder_values(case_id, party_id)
        merged_custom: dict[str, str] = custom_values or {}

        preview_items: list[FillPreviewItem] = []
        for mapping in mappings:
            fill_value: str = ""
            value_source: str = "empty"

            if mapping.semantic_label in placeholder_values:
                fill_value = str(placeholder_values[mapping.semantic_label])
                value_source = "auto"
            elif mapping.semantic_label in merged_custom:
                fill_value = merged_custom[mapping.semantic_label]
                value_source = "manual"

            preview_items.append(
                FillPreviewItem(
                    position_description=mapping.position_description,
                    semantic_label=mapping.semantic_label,
                    fill_value=fill_value,
                    value_source=value_source,
                    fill_type=mapping.fill_type,
                    mapping_id=mapping.id,
                )
            )

        logger.info(
            "填充预览生成: template_id=%d, case_id=%d, items=%d",
            template_id,
            case_id,
            len(preview_items),
        )
        return preview_items

    def get_custom_fields(self, template_id: int) -> list[dict[str, Any]]:
        """获取所有字段映射列表（供填充页面展示）"""
        from apps.documents.models.external_template import ExternalTemplateFieldMapping

        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
        ).order_by("sort_order", "id")

        fields: list[dict[str, Any]] = []
        for mapping in mappings:
            fields.append(
                {
                    "mapping_id": mapping.id,
                    "semantic_label": mapping.semantic_label,
                    "fill_type": mapping.fill_type,
                    "position_description": mapping.position_description,
                }
            )

        logger.info(
            "字段列表获取: template_id=%d, count=%d",
            template_id,
            len(fields),
        )
        return fields

    # ------------------------------------------------------------------
    # 内部辅助
    # ------------------------------------------------------------------

    def _get_placeholder_values(self, case_id: int, party_id: int | None = None) -> dict[str, str]:
        """
        从占位符体系获取案件+当事人的所有占位符值。

        1. 构建 context_data（case_id, party_id）
        2. 遍历 registry 中所有服务，调用 generate(context_data)
        3. 合并所有结果为 dict[str, str]
        """
        context_data: dict[str, Any] = {"case_id": case_id}
        if party_id is not None:
            context_data["party_id"] = party_id

        all_values: dict[str, str] = {}
        services = self._placeholder_registry.get_all_services()

        for service in services:
            try:
                result: dict[str, Any] = service.generate(context_data)
                for key, value in result.items():
                    all_values[key] = str(value) if value is not None else ""
            except Exception:
                logger.exception(
                    "占位符服务 %s 生成失败: case_id=%d",
                    service.name,
                    case_id,
                )

        logger.info(
            "占位符值获取: case_id=%d, party_id=%s, keys=%d",
            case_id,
            party_id,
            len(all_values),
        )
        return all_values

    # ------------------------------------------------------------------
    # 单次填充
    # ------------------------------------------------------------------

    def fill_template(
        self,
        template_id: int,
        case_id: int,
        party_id: int | None = None,
        custom_values: dict[str, str] | None = None,
        filled_by: Any = None,
    ) -> Any:
        """
        执行单次填充：
        1. 获取占位符值 + 自定义值
        2. 打开模板 .docx 副本
        3. 按 FieldMapping 逐一写入值
        4. 保存生成文件
        5. 创建 FillRecord
        6. 返回 FillRecord

        Requirements: 6.1, 6.2, 6.3, 6.4, 6.5, 6.6, 6.7, 6.8, 6.9, 6.10,
                      10.1, 10.2, 10.3, 10.4, 10.5, 10.6, 11.3, 15.3, 15.4,
                      15.6, 16.3, 16.4, 16.6
        """
        from docx import Document

        from apps.documents.models.external_template import ExternalTemplate, ExternalTemplateFieldMapping
        from apps.documents.models.fill_record import FillRecord

        template: ExternalTemplate = ExternalTemplate.objects.get(id=template_id)
        mappings = ExternalTemplateFieldMapping.objects.filter(
            template_id=template_id,
        ).order_by("sort_order", "id")

        # 1. 获取占位符值 + 合并自定义值
        placeholder_values: dict[str, str] = self._get_placeholder_values(case_id, party_id)
        merged_custom: dict[str, str] = custom_values or {}

        # 2. 打开模板 .docx 副本
        template_path: Path = Path(settings.MEDIA_ROOT) / template.file_path
        doc: Document = Document(str(template_path))

        # 3. 逐一写入
        filled_count: int = 0
        skipped_count: int = 0
        manual_needed: list[str] = []
        errors: list[str] = []

        for mapping in mappings:
            value: str = ""
            # 确定值来源：优先占位符体系，其次自定义值
            if mapping.semantic_label in placeholder_values:
                value = placeholder_values[mapping.semantic_label]
            elif mapping.semantic_label in merged_custom:
                value = merged_custom[mapping.semantic_label]
            else:
                manual_needed.append(mapping.semantic_label)
                skipped_count += 1
                logger.info(
                    "跳过无值字段: template_id=%d, label=%s",
                    template_id,
                    mapping.semantic_label,
                )
                continue

            # 根据 fill_type 调用对应写入方法
            success: bool = False
            try:
                if mapping.fill_type == "text":
                    success = self._write_text(doc, mapping.position_locator, value)
                elif mapping.fill_type == "checkbox":
                    success = self._write_checkbox(doc, mapping.position_locator, value)
                elif mapping.fill_type == "delete_inapplicable":
                    success = self._write_delete_inapplicable(doc, mapping.position_locator, value)
                else:
                    logger.warning(
                        "未知填充类型: template_id=%d, fill_type=%s",
                        template_id,
                        mapping.fill_type,
                    )
                    skipped_count += 1
                    continue
            except Exception:
                logger.exception(
                    "填充写入失败: template_id=%d, locator=%s",
                    template_id,
                    mapping.position_locator,
                )
                errors.append(str(_("位置 %(label)s 写入失败") % {"label": mapping.semantic_label}))
                continue

            if success:
                filled_count += 1
            else:
                skipped_count += 1
                errors.append(str(_("位置 %(label)s 写入未成功") % {"label": mapping.semantic_label}))

        # 4. 保存生成文件
        output_dir: Path = Path(settings.MEDIA_ROOT) / "documents" / "external_filled" / str(case_id)
        output_dir.mkdir(parents=True, exist_ok=True)
        output_uuid: str = str(uuid.uuid4())
        output_abs: Path = output_dir / f"{output_uuid}.docx"
        doc.save(str(output_abs))

        output_relative: str = str(Path("documents") / "external_filled" / str(case_id) / f"{output_uuid}.docx")

        # 5. 生成输出文件名
        party_name: str | None = None
        if party_id is not None:
            try:
                case_party_model = apps.get_model("cases", "CaseParty")
                party_obj = case_party_model.objects.get(id=party_id)
                party_name = str(party_obj)
            except Exception:
                logger.warning("获取当事人名称失败: party_id=%d", party_id)

        output_name: str = self._generate_output_filename(template.name, party_name)
        # 6. 构建填充报告
        report: dict[str, Any] = {
            "total_fields": len(list(mappings)),
            "filled_count": filled_count,
            "skipped_count": skipped_count,
            "manual_needed": manual_needed,
            "errors": errors,
        }

        # 7. 创建 FillRecord
        record: FillRecord = FillRecord.objects.create(
            case_id=case_id,
            template=template,
            party_id=party_id,
            filled_by=filled_by,
            file_path=output_relative,
            original_output_name=output_name,
            report_json=report,
            custom_values=merged_custom,
        )

        logger.info(
            "填充完成: record_id=%d, template_id=%d, case_id=%d, filled=%d, skipped=%d, errors=%d",
            record.id,
            template_id,
            case_id,
            filled_count,
            skipped_count,
            len(errors),
        )
        return record

    # ------------------------------------------------------------------
    # 写入方法
    # ------------------------------------------------------------------

    def _write_text(self, doc: Any, locator: dict[str, Any], value: str) -> bool:
        """
        写入文本值，保留原有格式属性。

        支持两种定位方式：
        - paragraph: 通过 paragraph_index 定位段落
        - table_cell: 通过 table_index + row + col 定位表格单元格

        Requirements: 6.2, 6.3, 10.1, 10.2, 10.3
        """
        locator_type: str = locator.get("type", "")

        try:
            if locator_type == "paragraph":
                para_index: int = locator.get("paragraph_index", 0)
                if para_index >= len(doc.paragraphs):
                    logger.warning(
                        "段落索引越界: index=%d, total=%d",
                        para_index,
                        len(doc.paragraphs),
                    )
                    return False

                paragraph = doc.paragraphs[para_index]
                runs = paragraph.runs
                if not runs:
                    # 无 run 时直接添加
                    paragraph.add_run(value)
                    return True

                # 保留第一个 run 的格式，替换文本
                first_run = runs[0]
                first_run.text = value
                # 清除后续 run 的文本
                for run in runs[1:]:
                    run.text = ""
                return True

            elif locator_type == "table_cell":
                table_index: int = locator.get("table_index", 0)
                row: int = locator.get("row", 0)
                col: int = locator.get("col", 0)

                if table_index >= len(doc.tables):
                    logger.warning(
                        "表格索引越界: index=%d, total=%d",
                        table_index,
                        len(doc.tables),
                    )
                    return False

                table = doc.tables[table_index]
                if row >= len(table.rows) or col >= len(table.columns):
                    logger.warning(
                        "单元格索引越界: row=%d, col=%d, rows=%d, cols=%d",
                        row,
                        col,
                        len(table.rows),
                        len(table.columns),
                    )
                    return False

                cell = table.cell(row, col)
                # 写入单元格第一个段落
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]
                    runs = paragraph.runs
                    if not runs:
                        paragraph.add_run(value)
                    else:
                        runs[0].text = value
                        for run in runs[1:]:
                            run.text = ""
                else:
                    cell.text = value
                return True

            else:
                logger.warning("未知定位器类型: %s", locator_type)
                return False

        except Exception:
            logger.exception("写入文本失败: locator=%s", locator)
            return False

    def _write_checkbox(self, doc: Any, locator: dict[str, Any], value: str) -> bool:
        """
        设置复选框勾选状态。

        通过解析文档 XML 找到复选框控件并设置 checked 状态。
        value 为 "true"/"1" 时勾选，否则取消勾选。

        Requirements: 6.2
        """
        try:
            checkbox_index: int = locator.get("checkbox_index", 0)
            checked: bool = value.lower() in ("true", "1", "yes")

            # 解析文档 XML 查找复选框
            body_xml: str = doc.element.xml
            root: ET.Element = ET.fromstring(body_xml)

            # Word 复选框命名空间
            ns: dict[str, str] = {
                "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
                "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            }

            # 查找所有 w14:checkbox 元素
            checkboxes: list[ET.Element] = root.findall(".//w14:checkbox", ns)

            if checkbox_index >= len(checkboxes):
                # 尝试旧版复选框格式 (w:fldChar + w:ffData)
                ff_checkboxes: list[ET.Element] = root.findall(".//w:ffData/w:checkBox", ns)
                if checkbox_index >= len(ff_checkboxes):
                    logger.warning(
                        "复选框索引越界: index=%d, w14=%d, ff=%d",
                        checkbox_index,
                        len(checkboxes),
                        len(ff_checkboxes),
                    )
                    return False

                # 旧版复选框：设置 w:default 或 w:checked
                cb: ET.Element = ff_checkboxes[checkbox_index]
                checked_elem: ET.Element | None = cb.find("w:checked", ns)
                if checked_elem is None:
                    checked_elem = cb.find("w:default", ns)

                if checked_elem is not None:
                    checked_elem.set(
                        f"{{{ns['w']}}}val",
                        "1" if checked else "0",
                    )
                return True

            # 新版复选框 (w14:checkbox)
            cb_elem: ET.Element = checkboxes[checkbox_index]
            checked_state: ET.Element | None = cb_elem.find("w14:checked", ns)
            if checked_state is not None:
                checked_state.set(
                    f"{{{ns['w14']}}}val",
                    "1" if checked else "0",
                )
            return True

        except Exception:
            logger.exception("写入复选框失败: locator=%s", locator)
            return False

    def _write_delete_inapplicable(
        self,
        doc: Any,
        locator: dict[str, Any],
        value: str,
    ) -> bool:
        """
        删除不适用项：value 为要保留的选项文本，直接替换整个位置内容。
        """
        try:
            locator_type: str = locator.get("type", "")

            paragraph = None
            if locator_type == "paragraph":
                para_index: int = locator.get("paragraph_index", 0)
                if para_index >= len(doc.paragraphs):
                    return False
                paragraph = doc.paragraphs[para_index]

            elif locator_type == "table_cell":
                table_index: int = locator.get("table_index", 0)
                row: int = locator.get("row", 0)
                col: int = locator.get("col", 0)

                if table_index >= len(doc.tables):
                    return False
                table = doc.tables[table_index]
                if row >= len(table.rows) or col >= len(table.columns):
                    return False
                cell = table.cell(row, col)
                if cell.paragraphs:
                    paragraph = cell.paragraphs[0]

            elif locator_type == "delete_inapplicable":
                para_index = locator.get("paragraph_index", 0)
                if para_index < len(doc.paragraphs):
                    paragraph = doc.paragraphs[para_index]

            if paragraph is None:
                logger.warning("无法定位删除不适用项段落: locator=%s", locator)
                return False

            # 直接用保留项替换
            runs = paragraph.runs
            if runs:
                runs[0].text = value
                for run in runs[1:]:
                    run.text = ""
            else:
                paragraph.add_run(value)

            return True

        except Exception:
            logger.exception("删除不适用项失败: locator=%s", locator)
            return False

    # ------------------------------------------------------------------
    # 文件名生成
    # ------------------------------------------------------------------

    def _generate_output_filename(
        self,
        template_name: str,
        party_name: str | None = None,
    ) -> str:
        """
        生成输出文件名。
        格式：有当事人 "{template_name}_{party_name}.docx"，无当事人 "{template_name}.docx"
        """
        base_name: str = template_name
        if party_name:
            base_name = f"{template_name}_{party_name}"
        return f"{base_name}.docx"

    # ------------------------------------------------------------------
    # 批量填充
    # ------------------------------------------------------------------

    def batch_fill(
        self,
        case_id: int,
        template_ids: list[int],
        party_ids: list[int] | None = None,
        custom_values: dict[str, dict[str, str]] | None = None,
        filled_by: Any = None,
    ) -> Any:
        """
        批量填充：
        1. 创建 BatchFillTask
        2. 遍历 template_ids × party_ids 组合
        3. 逐一调用 fill_template
        4. 失败的模板跳过，记录错误
        5. 所有文件打包 ZIP
        6. 生成汇总报告
        7. 更新 BatchFillTask

        Requirements: 5.8, 14.1, 14.2, 14.3, 14.4, 14.5, 14.6,
                      15.1, 15.2, 15.5, 16.5
        """
        from apps.documents.models.fill_record import BatchFillTask

        # 1. 创建 BatchFillTask
        batch_task: BatchFillTask = BatchFillTask.objects.create(
            case_id=case_id,
            initiated_by=filled_by,
        )
        batch_task.templates.set(template_ids)

        # 2. 遍历 template_ids × party_ids 组合
        effective_party_ids: list[int | None] = list(party_ids) if party_ids else [None]
        records: list[Any] = []
        summary_results: list[dict[str, Any]] = []

        for template_id in template_ids:
            # 获取该模板的自定义值
            tpl_custom: dict[str, str] = {}
            if custom_values and str(template_id) in custom_values:
                tpl_custom = custom_values[str(template_id)]
            elif custom_values and template_id in custom_values:  # type: ignore[operator]
                tpl_custom = custom_values[template_id]  # type: ignore[index]

            for party_id in effective_party_ids:
                try:
                    record = self.fill_template(
                        template_id=template_id,
                        case_id=case_id,
                        party_id=party_id,
                        custom_values=tpl_custom,
                        filled_by=filled_by,
                    )
                    # 关联到 batch_task
                    record.batch_task = batch_task
                    record.save(update_fields=["batch_task"])
                    records.append(record)
                    summary_results.append(
                        {
                            "template_id": template_id,
                            "party_id": party_id,
                            "status": "success",
                            "record_id": record.id,
                            "filled_count": record.report_json.get("filled_count", 0),
                            "skipped_count": record.report_json.get("skipped_count", 0),
                        }
                    )
                except Exception:
                    logger.exception(
                        "批量填充失败: template_id=%d, party_id=%s, case_id=%d",
                        template_id,
                        party_id,
                        case_id,
                    )
                    summary_results.append(
                        {
                            "template_id": template_id,
                            "party_id": party_id,
                            "status": "failed",
                            "error": str(_("模板 %(tid)s 填充失败") % {"tid": template_id}),
                        }
                    )

        # 5. 打包 ZIP
        zip_path: str = ""
        if records:
            zip_path = self._pack_to_zip(records)

        # 6. 汇总报告
        success_count: int = sum(1 for r in summary_results if r["status"] == "success")
        failed_count: int = sum(1 for r in summary_results if r["status"] == "failed")
        summary: dict[str, Any] = {
            "total": len(summary_results),
            "success": success_count,
            "failed": failed_count,
            "details": summary_results,
        }

        # 7. 更新 BatchFillTask
        batch_task.finished_at = timezone.now()
        batch_task.zip_file_path = zip_path
        batch_task.summary_json = summary
        batch_task.save(update_fields=["finished_at", "zip_file_path", "summary_json"])

        logger.info(
            "批量填充完成: batch_task_id=%d, case_id=%d, success=%d, failed=%d",
            batch_task.id,
            case_id,
            success_count,
            failed_count,
        )
        return batch_task

    def _pack_to_zip(self, records: list[Any]) -> str:
        """
        将多个填充文件打包为 ZIP，返回 ZIP 文件相对路径。

        Requirements: 14.3
        """
        if not records:
            return ""

        first_record = records[0]
        case_id: int = first_record.case_id
        batch_task_id: int = first_record.batch_task_id

        zip_dir: Path = Path(settings.MEDIA_ROOT) / "documents" / "external_filled" / str(case_id)
        zip_dir.mkdir(parents=True, exist_ok=True)

        zip_filename: str = f"batch_{batch_task_id}.zip"
        zip_abs: Path = zip_dir / zip_filename

        with zipfile.ZipFile(str(zip_abs), "w", zipfile.ZIP_DEFLATED) as zf:
            for record in records:
                file_abs: Path = Path(settings.MEDIA_ROOT) / record.file_path
                if file_abs.exists():
                    zf.write(str(file_abs), record.original_output_name)
                else:
                    logger.warning(
                        "打包 ZIP 时文件不存在: record_id=%d, path=%s",
                        record.id,
                        record.file_path,
                    )

        zip_relative: str = str(Path("documents") / "external_filled" / str(case_id) / zip_filename)

        logger.info(
            "ZIP 打包完成: path=%s, files=%d",
            zip_relative,
            len(records),
        )
        return zip_relative

    # ------------------------------------------------------------------
    # 历史查询
    # ------------------------------------------------------------------

    def get_fill_history_by_case(self, case_id: int) -> QuerySet[Any]:
        """
        按案件查询填充历史。

        Requirements: 18.3
        """
        from apps.documents.models.fill_record import FillRecord

        return (
            FillRecord.objects.filter(case_id=case_id)
            .select_related("template", "party", "filled_by", "batch_task")
            .order_by("-filled_at")
        )

    def get_fill_history_by_template(self, template_id: int) -> QuerySet[Any]:
        """
        按模板查询填充历史。

        Requirements: 18.4
        """
        from apps.documents.models.fill_record import FillRecord

        return (
            FillRecord.objects.filter(template_id=template_id)
            .select_related("case", "party", "filled_by", "batch_task")
            .order_by("-filled_at")
        )

    # ------------------------------------------------------------------
    # 重新填充
    # ------------------------------------------------------------------

    def re_fill(self, record_id: int, filled_by: Any = None) -> Any:
        """
        基于历史记录重新填充（相同案件+模板+当事人组合）。

        Requirements: 18.6
        """
        from apps.documents.models.fill_record import FillRecord

        old_record: FillRecord = FillRecord.objects.get(id=record_id)

        new_record = self.fill_template(
            template_id=old_record.template_id,
            case_id=old_record.case_id,
            party_id=old_record.party_id,
            custom_values=old_record.custom_values or {},
            filled_by=filled_by,
        )

        logger.info(
            "重新填充完成: old_record_id=%d, new_record_id=%d",
            record_id,
            new_record.id,
        )
        return new_record

    # ------------------------------------------------------------------
    # 文件可用性检查
    # ------------------------------------------------------------------

    def check_file_availability(self, record_id: int) -> bool:
        """
        检查历史填充文件是否可用，并更新 file_available 字段。

        Requirements: 18.7
        """
        from apps.documents.models.fill_record import FillRecord

        record: FillRecord = FillRecord.objects.get(id=record_id)
        file_abs: Path = Path(settings.MEDIA_ROOT) / record.file_path
        available: bool = file_abs.exists()

        if record.file_available != available:
            record.file_available = available
            record.save(update_fields=["file_available"])
            logger.info(
                "文件可用性更新: record_id=%d, available=%s",
                record_id,
                available,
            )

        return available

    # ------------------------------------------------------------------
    # 自定义值缓存
    # ------------------------------------------------------------------

    def save_custom_values(
        self,
        case_id: int,
        template_id: int,
        custom_values: dict[str, str],
    ) -> None:
        """
        保存自定义值作为案件+模板组合的默认值。

        将自定义值保存到最新的 FillRecord 中，若无记录则创建占位记录。

        Requirements: 16.5
        """
        from apps.documents.models.fill_record import FillRecord

        record = (
            FillRecord.objects.filter(
                case_id=case_id,
                template_id=template_id,
            )
            .order_by("-filled_at")
            .first()
        )

        if record is not None:
            record.custom_values = custom_values
            record.save(update_fields=["custom_values"])
        else:
            # 创建占位记录
            FillRecord.objects.create(
                case_id=case_id,
                template_id=template_id,
                file_path="",
                original_output_name="",
                custom_values=custom_values,
                file_available=False,
            )

        logger.info(
            "自定义值保存: case_id=%d, template_id=%d, keys=%d",
            case_id,
            template_id,
            len(custom_values),
        )

    def load_custom_values(
        self,
        case_id: int,
        template_id: int,
    ) -> dict[str, str]:
        """
        加载案件+模板组合的已保存自定义值。

        Requirements: 16.5
        """
        from apps.documents.models.fill_record import FillRecord

        record = (
            FillRecord.objects.filter(
                case_id=case_id,
                template_id=template_id,
            )
            .order_by("-filled_at")
            .first()
        )

        if record is not None and record.custom_values:
            result: dict[str, str] = dict(record.custom_values)
            logger.info(
                "自定义值加载: case_id=%d, template_id=%d, keys=%d",
                case_id,
                template_id,
                len(result),
            )
            return result

        return {}
