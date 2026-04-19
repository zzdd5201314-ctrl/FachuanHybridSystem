"""DOCX 导出服务 —— 负责将截图列表渲染为 Word 文档字节流。"""

from __future__ import annotations

import contextlib
import io
import logging
from collections.abc import Callable
from typing import Any

from django.core.files.base import ContentFile

from apps.chat_records.models import ChatRecordProject, ChatRecordScreenshot
from apps.core.exceptions import ValidationException

from .export_types import ExportLayout

logger = logging.getLogger(__name__)


class DocxExportService:
    """单一职责：生成 DOCX 导出文件。"""

    def export_docx(
        self,
        *,
        project: ChatRecordProject,
        screenshots: list[ChatRecordScreenshot],
        layout: ExportLayout,
        filename: str,
        progress_callback: Callable[[int, int, str], None] | None = None,
    ) -> ContentFile[bytes]:
        content = self._build_docx_bytes(
            project=project,
            screenshots=screenshots,
            layout=layout,
            progress_callback=progress_callback,
        )
        return ContentFile(content, name=filename)

    def _build_docx_bytes(
        self,
        *,
        project: ChatRecordProject,
        screenshots: list[ChatRecordScreenshot],
        layout: ExportLayout,
        progress_callback: Callable[[int, int, str], None] | None = None,
    ) -> bytes:
        if not screenshots:
            raise ValidationException("没有截图,无法导出")

        from docx import Document
        from docx.shared import Mm

        document = Document()
        self._setup_docx_sections(document, layout)

        images_per_page = layout.images_per_page
        cols = 1 if images_per_page == 1 else 2
        processed = 0
        total = len(screenshots)
        inserted_images = 0

        for idx in range(0, len(screenshots), images_per_page):
            batch = screenshots[idx : idx + images_per_page]
            table = document.add_table(rows=1, cols=cols)
            table.autofit = True
            row = table.rows[0]

            for col, shot in enumerate(batch):
                cell = row.cells[col]
                width_mm = 170 if cols == 1 else 80
                self._insert_docx_image(cell, shot, Mm(width_mm))
                inserted_images += 1
                processed += 1
                if progress_callback:
                    progress_callback(processed, total, "生成中")

            if idx + images_per_page < len(screenshots):
                document.add_page_break()

        if inserted_images == 0:
            raise ValidationException("Word 导出失败:未插入任何图片")

        buf = io.BytesIO()
        document.save(buf)
        return buf.getvalue()

    def _setup_docx_sections(self, document: Any, layout: ExportLayout) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm

        for section in document.sections:
            section.page_width = Mm(210)
            section.page_height = Mm(297)
            section.left_margin = Mm(15)
            section.right_margin = Mm(15)
            section.top_margin = Mm(15)
            section.bottom_margin = Mm(15)

        if layout.show_page_number:
            footer_para = document.sections[0].footer.paragraphs[0]
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer_para.add_run("第 ")
            self._add_field(footer_para, "PAGE")
            footer_para.add_run(" / ")
            self._add_field(footer_para, "NUMPAGES")
            footer_para.add_run(" 页")

        if layout.header_text:
            p = document.add_paragraph(layout.header_text)
            with contextlib.suppress(Exception):
                p.style = document.styles["Title"]

    def _insert_docx_image(self, cell: Any, shot: ChatRecordScreenshot, width: Any) -> None:
        from PIL import Image

        try:
            run = cell.paragraphs[0].add_run()
            with shot.image.open("rb") as fh, Image.open(fh) as img:
                rgb_img = img.convert("RGB")
                img_buf = io.BytesIO()
                rgb_img.save(img_buf, format="JPEG", quality=82, optimize=True)
                img_buf.seek(0)
                run.add_picture(img_buf, width=width)
        except Exception:
            raise ValidationException("Word 导出插图失败") from None

        title = (shot.title or "").strip()
        note = (shot.note or "").strip()
        if title:
            cell.add_paragraph(title)
        if note:
            cell.add_paragraph(note)

    def _add_field(self, paragraph: Any, field: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        r = run._r

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")

        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = f" {field} "

        fld_char_separate = OxmlElement("w:fldChar")
        fld_char_separate.set(qn("w:fldCharType"), "separate")

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        r.append(fld_char_begin)
        r.append(instr_text)
        r.append(fld_char_separate)
        r.append(fld_char_end)
