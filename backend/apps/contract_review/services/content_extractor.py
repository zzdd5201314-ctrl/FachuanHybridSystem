from __future__ import annotations

import logging
from dataclasses import dataclass
from pathlib import Path

from docx import Document

from apps.contract_review.services.exceptions import ExtractionError

logger = logging.getLogger(__name__)


@dataclass
class ExtractionResult:
    """提取结果，包含段落文本和索引映射"""

    paragraphs: list[str]
    index_map: list[int]  # index_map[filtered_idx] = original_idx in doc.paragraphs


class ContentExtractor:
    """合同内容提取器"""

    def extract_paragraphs(self, file_path: Path) -> list[str]:
        """从 docx 提取全部非空段落文本（向后兼容）"""
        result = self.extract_with_mapping(file_path)
        return result.paragraphs

    def extract_with_mapping(self, file_path: Path) -> ExtractionResult:
        """从 docx 提取段落文本，同时返回到原始 doc.paragraphs 的索引映射"""
        try:
            doc = Document(str(file_path))
        except Exception as e:
            raise ExtractionError(f"无法打开文件: {e}") from e

        paragraphs: list[str] = []
        index_map: list[int] = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                paragraphs.append(text)
                index_map.append(i)

        if not paragraphs:
            raise ExtractionError("文档内容为空，无法提取有效文本")

        logger.info("从 %s 提取到 %d 个段落", file_path.name, len(paragraphs))
        return ExtractionResult(paragraphs=paragraphs, index_map=index_map)
