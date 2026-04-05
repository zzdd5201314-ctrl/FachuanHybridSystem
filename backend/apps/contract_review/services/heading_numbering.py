from __future__ import annotations

import json
import logging
import re

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from apps.core.llm.service import LLMService

logger = logging.getLogger(__name__)


class HeadingNumbering:
    """通过 LLM 识别标题层级并设置 OOXML 多级自动编号"""

    def __init__(self, llm_service: LLMService | None = None) -> None:
        self._llm = llm_service

    def apply_numbering(self, doc: Document, model_name: str = "") -> None:
        """识别标题段落，定义并应用多级列表编号"""
        if not self._llm:
            logger.warning("未提供 LLM 服务，跳过标题编号")
            return

        headings = self._identify_headings_via_llm(doc, model_name)
        if not headings:
            logger.info("未识别到标题段落，跳过编号")
            return

        # 补充 LLM 漏识别的编号段落（原始有 numPr 且以编号前缀开头）
        headings = self._supplement_missed_headings(doc, headings)

        abstract_id = self._create_abstract_num(doc)

        # 按附件分隔符(-1)拆分为多个编号区域，每个区域独立 numId
        heading_indices = {idx for idx, _ in headings}
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        # 仅对 LLM 识别的标题段落去掉手动编号前缀
        self._strip_manual_numbers(doc, [(idx, lvl) for idx, lvl in headings if lvl >= 0])

        # 清除非标题段落的残留 numPr（numId=0 的无效编号引用）
        for i, p in enumerate(doc.paragraphs):
            if i in heading_indices:
                continue
            num_pr = p._element.find(f"{{{ns}}}pPr/{{{ns}}}numPr")
            if num_pr is not None:
                num_pr.getparent().remove(num_pr)

        # 拆分区域并分配 numId
        numbering_elem = doc.part.numbering_part.element
        sections: list[list[tuple[int, int]]] = [[]]
        for idx, lvl in headings:
            if lvl == -1:
                sections.append([])  # 新区域
            else:
                sections[-1].append((idx, lvl))

        applied = 0
        for section in sections:
            if not section:
                continue
            num_id = self._create_num_ref(numbering_elem, abstract_id)
            self._apply_num_to_paragraphs(doc, section, num_id)
            applied += len(section)

        logger.info("已为 %d 个标题段落应用编号（%d 个编号区域）", applied, len([s for s in sections if s]))

    def _identify_headings_via_llm(self, doc: Document, model_name: str) -> list[tuple[int, int]]:
        """用 LLM 识别标题段落及层级，返回 (段落索引, 层级0/1/2)"""
        lines: list[str] = []
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            if text:
                lines.append(f"[{i}] {text}")
        if not lines:
            return []

        text_block = "\n".join(lines)
        prompt = (
            "你是合同文档结构分析专家。以下是一份合同的所有段落（前面是段落索引号）。\n"
            "请识别哪些段落是标题/条款编号行，并判断层级：\n"
            '- level -1: 附件标题（如 "附件1：…" "附件一：…" "附件2-…"），用于分隔编号区域\n'
            '- level 0: 一级标题（如 "一、产品名称" "第一条 …" "第一章 …"）\n'
            '- level 1: 二级标题（如 "1." "1、" "（一）" 开头的条款）\n'
            '- level 2: 三级标题（如 "（1）" "①" "1.1" 开头的子条款）\n\n'
            "注意：\n"
            '- 合同标题（如 "XX合同"）不算条款标题，不要包含\n'
            "- 当事人信息行不算标题\n"
            "- 签署栏、日期行不算标题\n"
            "- 附件内部的条款标题也要识别（level 0/1/2），附件标题本身用 level -1\n"
            "- 只返回确实是标题的段落\n\n"
            "仅返回 JSON 数组，格式：\n"
            '[{"index": 段落索引, "level": 层级}]\n'
            "如果没有标题，返回 []。不要返回任何其他内容。\n\n"
            f"{text_block}"
        )
        for attempt in range(2):
            try:
                resp = self._llm.complete(
                    prompt=prompt,
                    model=model_name or None,
                    temperature=0.1,
                    fallback=False,
                )
                return self._parse_llm_response(resp.content, len(doc.paragraphs))
            except Exception:
                if attempt == 0:
                    logger.warning("标题识别 LLM 调用失败，重试中...")
                    continue
                logger.exception("标题识别 LLM 调用失败（已重试）")
                return []
        return []

    @staticmethod
    def _supplement_missed_headings(doc: Document, headings: list[tuple[int, int]]) -> list[tuple[int, int]]:
        """补充 LLM 漏识别的编号段落：原始有 numPr 的段落（有真实编号或有编号前缀文本）"""
        heading_indices = {idx for idx, _ in headings}
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        prefix_re = re.compile(r"^(?:\d+[、．.]\s*|[（(]\d+[)）]\s*|\d+\)\s*)")

        added = []
        for i, p in enumerate(doc.paragraphs):
            if i in heading_indices:
                continue
            num_pr = p._element.find(f"{{{ns}}}pPr/{{{ns}}}numPr")
            if num_pr is None:
                continue
            num_id_elem = num_pr.find(qn("w:numId"))
            ilvl_elem = num_pr.find(qn("w:ilvl"))
            num_id = int(num_id_elem.get(qn("w:val"), "0")) if num_id_elem is not None else 0
            orig_ilvl = int(ilvl_elem.get(qn("w:val"), "0")) if ilvl_elem is not None else 0

            if num_id > 0:
                # 有真实自动编号，保留原始层级（限制在0-2）
                level = min(orig_ilvl, 2)
            elif prefix_re.match(p.text.strip()):
                # numId=0 但有编号前缀文本，推断层级
                parent_level = -1
                for idx, lvl in sorted(headings, key=lambda x: x[0]):
                    if idx < i and lvl >= 0:
                        parent_level = lvl
                level = min(parent_level + 1, 2)
            else:
                continue

            added.append((i, level))
            heading_indices.add(i)

        if added:
            logger.info("补充识别 %d 个编号段落", len(added))
            result = headings + added
            result.sort(key=lambda x: x[0])
            return result
        return headings

    @staticmethod
    def _parse_llm_response(text: str, para_count: int) -> list[tuple[int, int]]:
        text = text.strip()
        if "```" in text:
            start = text.find("[")
            end = text.rfind("]") + 1
            if start >= 0 and end > start:
                text = text[start:end]
        try:
            data = json.loads(text)
        except json.JSONDecodeError:
            logger.warning("标题识别 JSON 解析失败: %s", text[:200])
            return []
        if not isinstance(data, list):
            return []
        results: list[tuple[int, int]] = []
        for item in data:
            if not isinstance(item, dict):
                continue
            idx = item.get("index")
            lvl = item.get("level")
            if isinstance(idx, int) and isinstance(lvl, int) and 0 <= idx < para_count and -1 <= lvl <= 2:
                results.append((idx, lvl))
        return results

    @staticmethod
    def _strip_manual_numbers(doc: Document, headings: list[tuple[int, int]]) -> None:
        """去掉标题段落开头的手动编号前缀"""
        # 长模式优先，避免短模式误匹配
        pattern = re.compile(
            r"^("
            r"第[一二三四五六七八九十百\d]+[条章节][、．.\s]*"
            r"|[一二三四五六七八九十百]+[、．.]\s*"
            r"|（[一二三四五六七八九十百\d]+）\s*"
            r"|\([一二三四五六七八九十百\d]+\)\s*"
            r"|[①②③④⑤⑥⑦⑧⑨⑩]\s*"
            r"|\d+\.\d+[、．.\s]*"
            r"|\d+[、．.]\s*"
            r")"
        )
        for para_idx, _ in headings:
            para = doc.paragraphs[para_idx]
            if not para.runs:
                continue
            full_text = para.text
            m = pattern.match(full_text)
            if not m:
                continue
            prefix_len = len(m.group())
            # 从 runs 中移除前缀字符
            remaining = prefix_len
            for run in para.runs:
                if remaining <= 0:
                    break
                run_len = len(run.text)
                if run_len <= remaining:
                    run.text = ""
                    remaining -= run_len
                else:
                    run.text = run.text[remaining:].lstrip()
                    remaining = 0
            # 删除空 runs 并清理第一个非空 run 的前导空格
            for run in list(para.runs):
                if not run.text:
                    run._element.getparent().remove(run._element)
                else:
                    run.text = run.text.lstrip()
                    break

    def _create_abstract_num(self, doc: Document) -> int:
        """创建 abstractNum 定义，返回 abstractNumId"""
        numbering_part = doc.part.numbering_part
        numbering_elem = numbering_part.element

        existing_ids = {int(a.get(qn("w:abstractNumId"), 0)) for a in numbering_elem.findall(qn("w:abstractNum"))}
        abstract_id = max(existing_ids, default=-1) + 1

        abstract_num = OxmlElement("w:abstractNum")
        abstract_num.set(qn("w:abstractNumId"), str(abstract_id))

        levels = [
            {"ilvl": "0", "numFmt": "chineseCounting", "lvlText": "%1、", "start": "1"},
            {"ilvl": "1", "numFmt": "decimal", "lvlText": "%2.", "start": "1"},
            {"ilvl": "2", "numFmt": "decimal", "lvlText": "（%3）", "start": "1"},
        ]

        for lvl_def in levels:
            lvl = OxmlElement("w:lvl")
            lvl.set(qn("w:ilvl"), lvl_def["ilvl"])

            start = OxmlElement("w:start")
            start.set(qn("w:val"), lvl_def["start"])
            lvl.append(start)

            num_fmt = OxmlElement("w:numFmt")
            num_fmt.set(qn("w:val"), lvl_def["numFmt"])
            lvl.append(num_fmt)

            suff = OxmlElement("w:suff")
            suff.set(qn("w:val"), "nothing")
            lvl.append(suff)

            lvl_text = OxmlElement("w:lvlText")
            lvl_text.set(qn("w:val"), lvl_def["lvlText"])
            lvl.append(lvl_text)

            lvl_jc = OxmlElement("w:lvlJc")
            lvl_jc.set(qn("w:val"), "left")
            lvl.append(lvl_jc)

            abstract_num.append(lvl)

        first_num = numbering_elem.find(qn("w:num"))
        if first_num is not None:
            first_num.addprevious(abstract_num)
        else:
            numbering_elem.append(abstract_num)

        return abstract_id

    @staticmethod
    def _create_num_ref(numbering_elem: object, abstract_id: int) -> int:
        """创建 num 引用指向 abstractNum，返回 numId（每次调用生成独立编号序列）"""
        existing_num_ids = {
            int(n.get(qn("w:numId"), 0))
            for n in numbering_elem.findall(qn("w:num"))  # type: ignore[union-attr]
        }
        num_id = max(existing_num_ids, default=0) + 1

        num_elem = OxmlElement("w:num")
        num_elem.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num_elem.append(abstract_ref)
        numbering_elem.append(num_elem)  # type: ignore[union-attr]

        return num_id

    @staticmethod
    def _apply_num_to_paragraphs(doc: Document, headings: list[tuple[int, int]], num_id: int) -> None:
        """将编号样式应用到标题段落"""
        for para_idx, level in headings:
            para = doc.paragraphs[para_idx]
            p_pr = para._element.find(qn("w:pPr"))
            if p_pr is None:
                p_pr = OxmlElement("w:pPr")
                para._element.insert(0, p_pr)

            old_num_pr = p_pr.find(qn("w:numPr"))
            if old_num_pr is not None:
                p_pr.remove(old_num_pr)

            # 清除段落自身的缩进，让编号定义的缩进生效
            old_ind = p_pr.find(qn("w:ind"))
            if old_ind is not None:
                p_pr.remove(old_ind)

            # 显式设置 ind=0 覆盖样式继承的缩进
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "0")
            ind.set(qn("w:firstLine"), "0")
            p_pr.append(ind)

            num_pr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), str(level))
            num_pr.append(ilvl)

            num_id_elem = OxmlElement("w:numId")
            num_id_elem.set(qn("w:val"), str(num_id))
            num_pr.append(num_id_elem)

            p_pr.append(num_pr)
